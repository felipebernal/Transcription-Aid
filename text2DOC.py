# -*- coding: utf-8 -*-
"""
Created on Fri Jul 29 10:11:05 2016

@author: bernalarangofa
"""
from docx import Document
from docx.shared import Inches
from os import path

def main(pathfileTXT):

    ptr = open(pathfileTXT, "r")  # text file I need to convert
    
    
    (head,tail)=path.split(pathfileTXT)
    (head2,tail2)=path.splitext(tail)
    filename=head2
    
    document = Document()
    
    document.add_heading(filename, 0)
    
    p = document.add_paragraph(ptr.read())
    
    document.save(head+'/'+filename+'.docx')